#!/usr/bin/env python3
"""
Тестовый скрипт для создания Word документа с форматированием
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_test_document_with_formatting():
    """Создает тестовый Word документ с различными типами форматирования"""
    doc = Document()
    
    # Заголовок документа
    title = doc.add_heading('Тестовый документ с форматированием', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Подзаголовок
    doc.add_heading('Введение', level=1)
    
    # Обычный параграф
    doc.add_paragraph('Это тестовый Word документ, созданный для проверки функциональности просмотра документов с сохранением форматирования.')
    
    # Параграф с жирным текстом
    p = doc.add_paragraph('Этот текст содержит ')
    p.add_run('жирный текст').bold = True
    p.add_run(' и обычный текст.')
    
    # Параграф с курсивом
    p = doc.add_paragraph('Этот текст содержит ')
    p.add_run('курсивный текст').italic = True
    p.add_run(' и обычный текст.')
    
    # Параграф с подчеркиванием
    p = doc.add_paragraph('Этот текст содержит ')
    p.add_run('подчеркнутый текст').underline = True
    p.add_run(' и обычный текст.')
    
    # Параграф с зачеркиванием
    p = doc.add_paragraph('Этот текст содержит ')
    p.add_run('зачеркнутый текст').font.strike = True
    p.add_run(' и обычный текст.')
    
    # Комбинированное форматирование
    p = doc.add_paragraph('Комбинированное форматирование: ')
    p.add_run('жирный курсив').bold = True
    p.add_run(' и ')
    p.add_run('жирный подчеркнутый').bold = True
    p.add_run(' текст.')
    
    # Заголовок второго уровня
    doc.add_heading('Списки', level=2)
    
    # Маркированный список
    doc.add_paragraph('Маркированный список:', style='List Bullet')
    doc.add_paragraph('Первый элемент', style='List Bullet')
    doc.add_paragraph('Второй элемент', style='List Bullet')
    doc.add_paragraph('Третий элемент', style='List Bullet')
    
    # Нумерованный список
    doc.add_paragraph('Нумерованный список:', style='List Number')
    doc.add_paragraph('Первый элемент', style='List Number')
    doc.add_paragraph('Второй элемент', style='List Number')
    doc.add_paragraph('Третий элемент', style='List Number')
    
    # Заголовок для таблицы
    doc.add_heading('Таблица', level=2)
    
    # Создаем таблицу
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    # Заполняем заголовки таблицы
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Колонка 1'
    hdr_cells[1].text = 'Колонка 2'
    hdr_cells[2].text = 'Колонка 3'
    
    # Делаем заголовки жирными
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Заполняем данные таблицы
    row_cells = table.rows[1].cells
    row_cells[0].text = 'Данные 1'
    row_cells[1].text = 'Данные 2'
    row_cells[2].text = 'Данные 3'
    
    row_cells = table.rows[2].cells
    row_cells[0].text = 'Данные 4'
    row_cells[1].text = 'Данные 5'
    row_cells[2].text = 'Данные 6'
    
    row_cells = table.rows[3].cells
    row_cells[0].text = 'Данные 7'
    row_cells[1].text = 'Данные 8'
    row_cells[2].text = 'Данные 9'
    
    # Заголовок для заключения
    doc.add_heading('Заключение', level=1)
    
    # Заключительный параграф
    doc.add_paragraph('Этот документ демонстрирует различные типы форматирования, которые должны корректно отображаться в веб-интерфейсе:')
    
    # Список возможностей
    doc.add_paragraph('• Заголовки разных уровней', style='List Bullet')
    doc.add_paragraph('• Жирный, курсивный и подчеркнутый текст', style='List Bullet')
    doc.add_paragraph('• Зачеркнутый текст', style='List Bullet')
    doc.add_paragraph('• Маркированные и нумерованные списки', style='List Bullet')
    doc.add_paragraph('• Таблицы с форматированием', style='List Bullet')
    
    # Сохраняем документ
    test_file_path = 'test_document_with_formatting.docx'
    doc.save(test_file_path)
    print(f"Тестовый документ с форматированием создан: {test_file_path}")
    return test_file_path

if __name__ == "__main__":
    create_test_document_with_formatting() 