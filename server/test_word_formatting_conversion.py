#!/usr/bin/env python3
"""
Тестовый скрипт для проверки конвертации Word документов с форматированием
"""

import os
import sys
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Добавляем путь к модулям проекта
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from routes.content_routes import convert_docx_to_html_with_formatting

def create_test_document():
    """Создает тестовый Word документ с форматированием"""
    doc = Document()
    
    # Заголовок документа
    title = doc.add_heading('Тестовый документ с форматированием', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Подзаголовок
    doc.add_heading('Введение', level=1)
    
    # Обычный параграф
    doc.add_paragraph('Это тестовый Word документ, созданный для проверки функциональности просмотра документов с сохранением форматирования.')
    
    # Параграф с жирным текстом
    p = doc.add_paragraph('Этот текст содержит ')
    p.add_run('жирный текст').bold = True
    p.add_run(' и обычный текст.')
    
    # Параграф с курсивом
    p = doc.add_paragraph('Этот текст содержит ')
    p.add_run('курсивный текст').italic = True
    p.add_run(' и обычный текст.')
    
    # Параграф с подчеркиванием
    p = doc.add_paragraph('Этот текст содержит ')
    p.add_run('подчеркнутый текст').underline = True
    p.add_run(' и обычный текст.')
    
    # Параграф с зачеркиванием
    p = doc.add_paragraph('Этот текст содержит ')
    p.add_run('зачеркнутый текст').font.strike = True
    p.add_run(' и обычный текст.')
    
    # Комбинированное форматирование
    p = doc.add_paragraph('Комбинированное форматирование: ')
    p.add_run('жирный курсив').bold = True
    p.add_run(' и ')
    p.add_run('жирный подчеркнутый').bold = True
    p.add_run(' текст.')
    
    # Заголовок второго уровня
    doc.add_heading('Списки', level=2)
    
    # Маркированный список
    doc.add_paragraph('Маркированный список:', style='List Bullet')
    doc.add_paragraph('Первый элемент', style='List Bullet')
    doc.add_paragraph('Второй элемент', style='List Bullet')
    doc.add_paragraph('Третий элемент', style='List Bullet')
    
    # Нумерованный список
    doc.add_paragraph('Нумерованный список:', style='List Number')
    doc.add_paragraph('Первый элемент', style='List Number')
    doc.add_paragraph('Второй элемент', style='List Number')
    doc.add_paragraph('Третий элемент', style='List Number')
    
    # Заголовок для таблицы
    doc.add_heading('Таблица', level=2)
    
    # Создаем таблицу
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    # Заполняем заголовки таблицы
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Колонка 1'
    hdr_cells[1].text = 'Колонка 2'
    hdr_cells[2].text = 'Колонка 3'
    
    # Делаем заголовки жирными
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Заполняем данные таблицы
    row_cells = table.rows[1].cells
    row_cells[0].text = 'Данные 1'
    row_cells[1].text = 'Данные 2'
    row_cells[2].text = 'Данные 3'
    
    row_cells = table.rows[2].cells
    row_cells[0].text = 'Данные 4'
    row_cells[1].text = 'Данные 5'
    row_cells[2].text = 'Данные 6'
    
    row_cells = table.rows[3].cells
    row_cells[0].text = 'Данные 7'
    row_cells[1].text = 'Данные 8'
    row_cells[2].text = 'Данные 9'
    
    # Заголовок для заключения
    doc.add_heading('Заключение', level=1)
    
    # Заключительный параграф
    doc.add_paragraph('Этот документ демонстрирует различные типы форматирования, которые должны корректно отображаться в веб-интерфейсе:')
    
    # Список возможностей
    doc.add_paragraph('• Заголовки разных уровней', style='List Bullet')
    doc.add_paragraph('• Жирный, курсивный и подчеркнутый текст', style='List Bullet')
    doc.add_paragraph('• Зачеркнутый текст', style='List Bullet')
    doc.add_paragraph('• Маркированные и нумерованные списки', style='List Bullet')
    doc.add_paragraph('• Таблицы с форматированием', style='List Bullet')
    
    # Сохраняем документ
    test_file_path = 'test_document_with_formatting.docx'
    doc.save(test_file_path)
    print(f"Тестовый документ с форматированием создан: {test_file_path}")
    return test_file_path

def test_conversion(file_path):
    """Тестирует конвертацию Word документа в HTML"""
    try:
        print("\n=== Тестирование конвертации Word документа в HTML ===")
        
        # Конвертируем документ
        html_content = convert_docx_to_html_with_formatting(file_path)
        
        print("\n=== Результат конвертации ===")
        print(html_content)
        print("\n=== Конец результата ===\n")
        
        # Сохраняем результат в файл для просмотра
        output_file = 'test_conversion_result.html'
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(f"""
<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Результат конвертации</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            margin: 0;
            padding: 20px;
            background-color: #f8f9fa;
        }}
        .container {{
            max-width: 900px;
            margin: 0 auto;
            background-color: white;
            padding: 30px;
            border-radius: 8px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        .docx-heading {{
            color: #2c3e50;
            margin-top: 20px;
            margin-bottom: 10px;
            font-weight: 600;
        }}
        .docx-paragraph {{
            margin-bottom: 12px;
            text-align: justify;
        }}
        .docx-table {{
            width: 100%;
            margin: 15px 0;
            border-collapse: collapse;
            border: 1px solid #ddd;
        }}
        .docx-cell {{
            padding: 8px;
            border: 1px solid #ddd;
            vertical-align: top;
        }}
        .docx-list {{
            margin: 15px 0;
            padding-left: 20px;
        }}
        .docx-list-item {{
            margin-bottom: 5px;
            line-height: 1.6;
        }}
        strong {{
            font-weight: 600;
        }}
        em {{
            font-style: italic;
        }}
        u {{
            text-decoration: underline;
        }}
        del {{
            text-decoration: line-through;
            color: #6c757d;
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>Результат конвертации Word документа</h1>
        <div class="content">
            {html_content}
        </div>
    </div>
</body>
</html>
            """)
        
        print(f"Результат сохранен в файл: {output_file}")
        return True
        
    except Exception as e:
        print(f"Ошибка при тестировании конвертации: {e}")
        return False

if __name__ == "__main__":
    # Создаем тестовый документ
    test_file = create_test_document()
    
    # Тестируем конвертацию
    success = test_conversion(test_file)
    
    if success:
        print("✅ Тестирование прошло успешно!")
    else:
        print("❌ Тестирование завершилось с ошибкой!") 