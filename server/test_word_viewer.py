#!/usr/bin/env python3
"""
Тестовый скрипт для проверки функциональности просмотра Word документов
"""

import os
import sys
from docx import Document
import docx2txt

def create_test_document():
    """Создает тестовый Word документ"""
    doc = Document()
    
    # Добавляем заголовок
    doc.add_heading('Тестовый документ для проверки функциональности', 0)
    
    # Добавляем параграфы
    doc.add_paragraph('Это тестовый Word документ, созданный для проверки функциональности просмотра документов в веб-интерфейсе.')
    
    doc.add_paragraph('Второй абзац содержит дополнительную информацию о возможностях системы.')
    
    # Добавляем подзаголовок
    doc.add_heading('Функциональность', level=1)
    
    # Добавляем список
    doc.add_paragraph('Основные возможности:', style='List Bullet')
    doc.add_paragraph('Просмотр Word документов в браузере', style='List Bullet')
    doc.add_paragraph('Конвертация в HTML формат', style='List Bullet')
    doc.add_paragraph('Интеграция с RAG системой', style='List Bullet')
    
    # Добавляем еще один параграф
    doc.add_paragraph('Этот документ демонстрирует, как система обрабатывает различные элементы форматирования Word документов.')
    
    # Сохраняем документ
    test_file_path = 'test_document.docx'
    doc.save(test_file_path)
    print(f"Тестовый документ создан: {test_file_path}")
    return test_file_path

def test_docx2txt(file_path):
    """Тестирует извлечение текста из Word документа"""
    try:
        text = docx2txt.process(file_path)
        print("\n=== Извлеченный текст ===")
        print(text)
        print("=== Конец текста ===\n")
        return True
    except Exception as e:
        print(f"Ошибка при извлечении текста: {e}")
        return False

def main():
    """Основная функция тестирования"""
    print("Тестирование функциональности просмотра Word документов")
    print("=" * 60)
    
    # Создаем тестовый документ
    test_file = create_test_document()
    
    # Тестируем извлечение текста
    if test_docx2txt(test_file):
        print("✅ Тест извлечения текста прошел успешно")
    else:
        print("❌ Тест извлечения текста не прошел")
        return False
    
    # Проверяем размер файла
    file_size = os.path.getsize(test_file)
    print(f"📄 Размер тестового файла: {file_size} байт")
    
    # Удаляем тестовый файл
    try:
        os.remove(test_file)
        print(f"🗑️ Тестовый файл удален: {test_file}")
    except Exception as e:
        print(f"⚠️ Не удалось удалить тестовый файл: {e}")
    
    print("\n🎉 Все тесты завершены успешно!")
    return True

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1) 